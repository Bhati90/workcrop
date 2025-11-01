import os
import json
from docx import Document
import pandas as pd
import re

# --- CONFIGURATION ---
DOCX_FILES = [
    '35.docx',
    '36.docx',
    '15.docx',
    'cri.docx',

    'th.docx',
]

XLSX_FILES = [
    'all.xlsx',
    'pro.xlsx',
    'col.xlsx',
]

OUTPUT_FILE = 'knowledge_base.json'

# --- COMPANY INFO ---
COMPANY_INFO = {
    "name": "Agriculture Service Platform",
    "tagline": "Best platform for all your farming needs - सबसे अच्छा कृषि सेवा मंच",
    "services": [
        "Farm Labor Connection - मजूर कनेक्शन सर्विस - Connect farmers with verified workers for pruning (कटाई), harvesting (तोड), spraying (फवारणी)",
        "Crop Protection Schedule - फसल संरक्षण शेड्यूल - Complete spray schedules for all grape varieties",
        "Product Supply - उत्पाद आपूर्ति - Fertilizers, pesticides, micronutrients at best prices",
        "Agri-Education - कृषि शिक्षण - Expert guidance on farming practices",
        "Equipment & Tools - उपकरण - Farm equipment and machinery support",
        # "Transport Services - वाहतूक सेवा - Transportation for produce and materials",
        # "Storage Solutions - साठवण सुविधा - Cold storage and warehousing",
    ],
    "pricing_promise": "We provide the BEST rates in the market. No need to worry about pricing - we are the most trusted platform. Our team will contact you within 24 hours with personalized pricing for your specific needs.",
    "coverage": "We serve farmers across Maharashtra (Satara, Pune, Nashik, Sangli, etc.) and expanding to more areas",
    "languages": "We support Hindi (हिंदी), Marathi (मराठी), English, and Hinglish"
}

# --- GRAPE VARIETIES ---
GRAPE_VARIETIES = [
    {
        "name": "ARD 35 / Arra 35",
        "aliases": ["ard35", "ard 35", "arra35", "arra 35", "अरा ३५", "एआरडी ३५"],
        "description": "ARD 35 (Arra 35) is a premium black seeded grape variety popular in Maharashtra. Known for excellent yield and market demand.",
        "file": "35.docx"
    },
    {
        "name": "ARD 36 / Arra 36",
        "aliases": ["ard36", "ard 36", "arra36", "arra 36", "अरा ३६", "एआरडी ३६"],
        "description": "ARD 36 (Arra 36) is another premium grape variety with high export potential.",
        "file": "36.docx"
    },
    {
        "name": "Arra 15",
        "aliases": ["arra15", "arra 15", "अरा १५", "15"],
        "description": "Arra 15 is a popular grape variety in Maharashtra known for good sweetness.",
        "file": "15.docx"
    },
    {
        "name": "Crimson Seedless",
        "aliases": ["crimson", "crimpson", "क्रिमसन"],
        "description": "Crimson Seedless is a globally popular red seedless grape variety with high market value.",
        "file": "cri.docx"
    },
    {
        "name": "Thompson Seedless",
        "aliases": ["thompson", "thomson", "थॉम्पसन"],
        "description": "Thompson Seedless is one of the most widely cultivated seedless grape varieties worldwide.",
        "file": "th.docx"
    },
]

def create_company_chunks():
    """Create chunks from company information"""
    chunks = []
    
    # Main company info
    chunks.append({
        "source": "Company Info",
        "type": "company_overview",
        "content": f"{COMPANY_INFO['name']} - {COMPANY_INFO['tagline']}. We provide: {' | '.join(COMPANY_INFO['services'])}"
    })
    
    # Pricing promise
    chunks.append({
        "source": "Company Info",
        "type": "pricing_policy",
        "content": f"Pricing Policy: {COMPANY_INFO['pricing_promise']}"
    })
    
    # Services details
    for service in COMPANY_INFO['services']:
        chunks.append({
            "source": "Company Info",
            "type": "service_detail",
            "content": service
        })
    
    # Coverage
    chunks.append({
        "source": "Company Info",
        "type": "coverage_area",
        "content": f"Service Coverage: {COMPANY_INFO['coverage']}"
    })
    
    # Languages
    chunks.append({
        "source": "Company Info",
        "type": "language_support",
        "content": f"Language Support: {COMPANY_INFO['languages']}"
    })
    
    return chunks

def create_variety_chunks():
    """Create grape variety definition chunks"""
    chunks = []
    
    for variety in GRAPE_VARIETIES:
        # Main variety info
        chunks.append({
            "source": "Grape Varieties Database",
            "type": "variety_definition",
            "content": f"{variety['name']}: {variety['description']} Schedule available in {variety['file']}"
        })
        
        # Alias mappings
        for alias in variety['aliases']:
            chunks.append({
                "source": "Grape Varieties Database",
                "type": "variety_alias",
                "content": f"'{alias}' is {variety['name']}"
            })
    
    return chunks

def process_docx(filepath):
    """Extract text from DOCX files with smart chunking"""
    print(f"📄 Processing {os.path.basename(filepath)}...")
    
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"  ❌ Error: {e}")
        return []
    
    chunks = []
    source_name = os.path.basename(filepath).replace('.docx', '')
    
    # Extract paragraphs
    current_section = ""
    current_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            continue
        
        # Detect section headers (usually bold or short)
        if len(text) < 100 and (para.style.name.startswith('Heading') or text.isupper()):
            # Save previous section
            if current_content:
                chunks.append({
                    "source": source_name,
                    "type": "crop_schedule_section",
                    "content": f"{current_section}: {' '.join(current_content)}"
                })
            
            # Start new section
            current_section = text
            current_content = []
        else:
            current_content.append(text)
    
    # Add last section
    if current_content:
        chunks.append({
            "source": source_name,
            "type": "crop_schedule_section",
            "content": f"{current_section}: {' '.join(current_content)}"
        })
    
    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        try:
            # Get headers
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            
            # Process each row
            for row_idx, row in enumerate(table.rows[1:], 1):
                cells = [cell.text.strip() for cell in row.cells]
                
                # Combine headers with values
                row_data = []
                for header, cell in zip(headers, cells):
                    if cell and cell != '-':
                        row_data.append(f"{header}: {cell}")
                
                if row_data:
                    chunks.append({
                        "source": source_name,
                        "type": "crop_schedule_table",
                        "content": " | ".join(row_data)
                    })
        except Exception as e:
            print(f"  ⚠️ Table {table_idx} error: {e}")
    
    print(f"  ✅ Extracted {len(chunks)} chunks")
    return chunks

def process_xlsx(filepath):
    """Extract data from Excel files with smart chunking"""
    print(f"📊 Processing {os.path.basename(filepath)}...")
    
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(filepath)
        chunks = []
        source_name = os.path.basename(filepath).replace('.xlsx', '')
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(filepath, sheet_name=sheet_name)
            
            # Clean column names
            df.columns = df.columns.str.strip()
            
            # Process each row
            for idx, row in df.iterrows():
                row_data = []
                
                for col in df.columns:
                    value = row[col]
                    
                    # Skip empty values
                    if pd.isna(value) or str(value).strip() == '':
                        continue
                    
                    # Format value
                    if isinstance(value, (int, float)):
                        value = str(value)
                    
                    row_data.append(f"{col}: {value}")
                
                if row_data:
                    chunks.append({
                        "source": f"{source_name} - {sheet_name}",
                        "type": "product_data",
                        "content": " | ".join(row_data)
                    })
        
        print(f"  ✅ Extracted {len(chunks)} chunks")
        return chunks
        
    except Exception as e:
        print(f"  ❌ Error: {e}")
        return []

def clean_and_deduplicate(chunks):
    """Remove duplicates and clean content"""
    seen = set()
    cleaned = []
    
    for chunk in chunks:
        content = chunk['content'].strip()
        
        # Skip very short chunks
        if len(content) < 10:
            continue
        
        # Skip duplicates
        if content in seen:
            continue
        
        seen.add(content)
        cleaned.append(chunk)
    
    return cleaned

def main():
    print("=" * 60)
    print("🌾 BUILDING ENHANCED KNOWLEDGE BASE")
    print("=" * 60)
    
    all_chunks = []
    
    # 1. Add company info
    print("\n📋 Adding company information...")
    company_chunks = create_company_chunks()
    all_chunks.extend(company_chunks)
    print(f"  ✅ Added {len(company_chunks)} company chunks")
    
    # 2. Add variety definitions
    print("\n🍇 Adding grape variety definitions...")
    variety_chunks = create_variety_chunks()
    all_chunks.extend(variety_chunks)
    print(f"  ✅ Added {len(variety_chunks)} variety chunks")
    
    # 3. Process DOCX files
    print("\n📚 Processing DOCX files...")
    for filename in DOCX_FILES:
        if os.path.exists(filename):
            chunks = process_docx(filename)
            all_chunks.extend(chunks)
        else:
            print(f"  ⚠️ File not found: {filename}")
    
    # 4. Process XLSX files
    print("\n📊 Processing XLSX files...")
    for filename in XLSX_FILES:
        if os.path.exists(filename):
            chunks = process_xlsx(filename)
            all_chunks.extend(chunks)
        else:
            print(f"  ⚠️ File not found: {filename}")
    
    # 5. Clean and deduplicate
    print("\n🧹 Cleaning and deduplicating...")
    all_chunks = clean_and_deduplicate(all_chunks)
    
    # 6. Save to JSON
    with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
        json.dump(all_chunks, f, indent=2, ensure_ascii=False)
    
    print("\n" + "=" * 60)
    print(f"✅ SUCCESS!")
    print(f"📦 Created {OUTPUT_FILE} with {len(all_chunks)} chunks")
    print("=" * 60)
    
    # Show sample
    print("\n📝 Sample chunks:")
    for i, chunk in enumerate(all_chunks[:3], 1):
        print(f"\n{i}. [{chunk['type']}] from {chunk['source']}")
        print(f"   {chunk['content'][:100]}...")

if __name__ == "__main__":
    main()